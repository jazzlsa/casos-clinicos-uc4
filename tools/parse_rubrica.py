#!/usr/bin/env python3
"""Importa um arquivo de referencia e gera os dados que o jogo usa.

O site e estatico (GitHub Pages) e nao roda Python, entao a fonte do conteudo
entra OFFLINE: este script roda no PC e gera dois artefatos que o site le:

  1. rubrica/<nome>_extraido.txt  - transcricao fiel do arquivo (paragrafos +
     tabelas, marcando negrito e hierarquia). E com isso que o conteudo do
     checklist e montado com precisao.
  2. data/rubrica.json            - saida do checklist (categorias do jogo),
     ja com a estrutura; as categorias reais sao preenchidas a partir da
     transcricao.

Uso:
    venv/Scripts/python.exe tools/parse_rubrica.py rubrica/NOME_ARQUIVO
    (ou: python tools/parse_rubrica.py rubrica/NOME_ARQUIVO)

Dependencia: python-docx  ->  pip install python-docx
"""
import argparse
import json
import sys
from pathlib import Path

PROJETO = Path(__file__).resolve().parents[1]
RUBRICA_DIR = PROJETO / "rubrica"
DATA_DIR = PROJETO / "data"


def _linhas_paragrafos(doc):
    out = []
    for p in doc.paragraphs:
        texto = p.text.strip()
        if not texto:
            continue
        negrito = any(r.bold for r in p.runs if r.bold)
        prefixo = "▋ " if negrito else ""
        out.append(f"{prefixo}{texto}")
    return out


def _linhas_tabelas(doc):
    out = []
    for t_idx, tabela in enumerate(doc.tables, 1):
        out.append(f"\n[TABELA {t_idx}]")
        for linha in tabela.rows:
            celulas = [c.text.strip().replace("\n", " ⏎ ") for c in linha.cells]
            # remove colunas repetidas que o Word duplica na mesma linha
            dedup = []
            for c in celulas:
                if not dedup or dedup[-1] != c:
                    dedup.append(c)
            out.append(" | ".join(dedup))
    return out


def extrair(arquivo_path: Path) -> str:
    import docx  # importa aqui pra mensagem de erro ficar clara

    doc = docx.Document(str(arquivo_path))
    linhas = _linhas_paragrafos(doc) + _linhas_tabelas(doc)
    cab = f"ARQUIVO: {arquivo_path.name}\n{'=' * 60}"
    return cab + "\n" + "\n".join(linhas)


def escrever_json(scaffold, extraido_txt: Path):
    payload = {
        "_comentario": (
            "Checklist de avaliacao da anamnese (UC4). Gerado a partir de "
            f"{extraido_txt.name}. Cada categoria tem 'nome' e 'pontos_max'; "
            "o app soma por categoria e mostra nota/max."
        ),
        "categorias": scaffold,
    }
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    (DATA_DIR / "rubrica.json").write_text(
        json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8"
    )


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("arquivo", help="Caminho do arquivo de referencia a importar")
    args = ap.parse_args()

    arquivo_path = Path(args.arquivo)
    if not arquivo_path.exists():
        print(f"Arquivo nao encontrado: {arquivo_path}", file=sys.stderr)
        return 2

    try:
        texto = extrair(arquivo_path)
    except ImportError:
        print("Faltou a dependencia 'python-docx'.\nRode:  pip install python-docx", file=sys.stderr)
        return 3
    except Exception as e:  # noqa: BLE001
        print(f"Falha ao ler o arquivo: {e}", file=sys.stderr)
        return 1

    RUBRICA_DIR.mkdir(parents=True, exist_ok=True)
    extraido_txt = RUBRICA_DIR / f"{arquivo_path.stem}_extraido.txt"
    extraido_txt.write_text(texto, encoding="utf-8")

    escrever_json([], extraido_txt)

    print(f"[OK] Transcricao salva em:  {extraido_txt}")
    print(f"[OK] data/rubrica.json gerado com a estrutura (categorias vazias).")
    print()
    print("Proximo passo: o conteudo das categorias e preenchido lendo a")
    print("transcricao acima (envie o conteudo dela pra eu montar o checklist,")
    print("ou preencha data/rubrica.json manualmente no formato:")
    print('  {"categorias": [{"nome": "Comunicacao", "pontos_max": 5}, ...]}')
    return 0


if __name__ == "__main__":
    sys.exit(main())
